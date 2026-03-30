"""
Script: xuat_bao_cao.py
Mục đích: Tự động tạo file báo cáo Word (.docx)
Thư viện: pip install python-docx
Cách dùng: python xuat_bao_cao.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# =============================================
# DỮ LIỆU BÁO CÁO
# =============================================
DU_LIEU = {
    "tieu_de":    "BÁO CÁO TUẦN 13",
    "ky_bao_cao": "24/03/2025 – 28/03/2025",
    "nguoi_viet": "Nguyễn Văn A",
    "gui_den":    "Trưởng nhóm Phát triển",
    "ngay_viet":  date.today().strftime("%d/%m/%Y"),

    "tom_tat": (
        "Tuần này hoàn thành 80% kế hoạch (8/10 task). "
        "Tính năng đăng nhập đã xong và chạy ổn định trên staging. "
        "Module thanh toán chậm 2 ngày do lỗi API đối tác — đang chờ phản hồi. "
        "Đã họp kickoff dự án B với khách hàng thành công."
    ),

    "ket_qua": [
        "Hoàn thành tính năng đăng nhập + đăng ký (100%) — pass 47/47 test case",
        "Fix 12 bug tồn đọng từ sprint trước — giảm từ 20 xuống còn 8 bug",
        "Họp kickoff dự án B — xác nhận 15 yêu cầu, có biên bản đầy đủ",
        "Viết tài liệu kỹ thuật module xác thực — 12 trang, đã review xong",
    ],

    "van_de": [
        "API thanh toán lỗi 500 khi gọi /charge — chưa rõ nguyên nhân từ vendor",
        "Đã gửi ticket #4521, dự kiến phản hồi thứ 2 tuần tới",
        "Ảnh hưởng: module thanh toán chậm 2 ngày so với kế hoạch",
    ],

    "ke_hoach": [
        "Hoàn thiện module thanh toán (ưu tiên cao nhất) — Deadline: 02/04",
        "Bắt đầu thiết kế màn hình dashboard — Deadline: 04/04",
        "Review code 2 thành viên mới (Minh + Hoa)",
        "Cập nhật tài liệu API sau khi module thanh toán xong",
    ],

    "de_xuat": [
        "Cần tài khoản sandbox cổng thanh toán ABC để tự debug — tiết kiệm 1-2 ngày",
        "Đề xuất họp 15 phút vào thứ 2 hàng tuần để đồng bộ tiến độ cả nhóm",
    ],
}

# =============================================
# MÀU SẮC
# =============================================
MAU_TIEU_DE  = RGBColor(0x1F, 0x4E, 0x79)
MAU_XAM      = RGBColor(0x55, 0x55, 0x55)

PHAN = [
    ("ket_qua",  "✅  KẾT QUẢ ĐẠT ĐƯỢC",         "1B7A3E", "EBF3FB"),
    ("van_de",   "⚠️  VẤN ĐỀ / KHÓ KHĂN",         "B86A00", "FFF3CD"),
    ("ke_hoach", "📅  KẾ HOẠCH KỲ TỚI",            "1F6B4E", "E8F5E9"),
    ("de_xuat",  "🙋  ĐỀ XUẤT / YÊU CẦU HỖ TRỢ",  "6A1B9A", "F3E5F5"),
]

# =============================================
# HÀM TIỆN ÍCH
# =============================================
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_padding(cell, top=80, bottom=80, left=140, right=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def them_doan(doc, text, dam=False, nghieng=False, co=12,
              mau=None, can_le=WD_ALIGN_PARAGRAPH.LEFT,
              truoc=0, sau=6):
    p = doc.add_paragraph()
    p.alignment = can_le
    p.paragraph_format.space_before = Pt(truoc)
    p.paragraph_format.space_after  = Pt(sau)
    run = p.add_run(text)
    run.bold        = dam
    run.italic      = nghieng
    run.font.size   = Pt(co)
    run.font.name   = "Arial"
    if mau:
        run.font.color.rgb = mau
    return p


def them_tieu_de_phan(doc, bieu_tuong_va_ten, hex_color: str):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)

    # Kẻ đường dưới theo màu phần
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)

    run = p.add_run(bieu_tuong_va_ten)
    run.bold            = True
    run.font.size       = Pt(13)
    run.font.name       = "Arial"
    run.font.color.rgb  = RGBColor.from_string(hex_color)
    return p


def them_bang(doc, danh_sach: list, mau_nen: str):
    table = doc.add_table(rows=len(danh_sach), cols=1)
    table.style = "Table Grid"

    # Độ rộng toàn trang
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "9072")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    for i, muc in enumerate(danh_sach):
        cell = table.rows[i].cells[0]
        set_cell_bg(cell, mau_nen)
        set_cell_padding(cell)
        p    = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run  = p.add_run(f"  •  {muc}")
        run.font.size = Pt(11)
        run.font.name = "Arial"

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# =============================================
# TẠO TÀI LIỆU
# =============================================
def tao_bao_cao(du_lieu: dict, ten_file: str = "bao-cao.docx"):
    doc = Document()

    # Trang A4, căn lề 2.5cm
    for section in doc.sections:
        section.page_height     = Cm(29.7)
        section.page_width      = Cm(21.0)
        section.left_margin     = Cm(2.5)
        section.right_margin    = Cm(2.5)
        section.top_margin      = Cm(2.0)
        section.bottom_margin   = Cm(2.0)

    # Tiêu đề
    them_doan(doc, du_lieu["tieu_de"],
              dam=True, co=20, mau=MAU_TIEU_DE,
              can_le=WD_ALIGN_PARAGRAPH.CENTER, sau=4)

    them_doan(doc, f"Kỳ báo cáo: {du_lieu['ky_bao_cao']}",
              co=10, mau=MAU_XAM,
              can_le=WD_ALIGN_PARAGRAPH.CENTER, sau=2)

    them_doan(doc,
              f"Người viết: {du_lieu['nguoi_viet']}   |   "
              f"Gửi đến: {du_lieu['gui_den']}   |   "
              f"Ngày: {du_lieu['ngay_viet']}",
              co=9, mau=MAU_XAM,
              can_le=WD_ALIGN_PARAGRAPH.CENTER, sau=10)

    # Tóm tắt
    them_tieu_de_phan(doc, "📌  TÓM TẮT", "1F4E79")
    them_doan(doc, du_lieu["tom_tat"], nghieng=True, co=11, truoc=4, sau=6)

    # Các phần nội dung
    for key, ten_phan, mau_hex, mau_nen in PHAN:
        them_tieu_de_phan(doc, ten_phan, mau_hex)
        them_bang(doc, du_lieu[key], mau_nen)

    doc.save(ten_file)
    print(f"✅ Đã tạo file: {ten_file}")


if __name__ == "__main__":
    tao_bao_cao(DU_LIEU, "bao-cao-tuan.docx")
